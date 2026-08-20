"""Bounded local document ingestion and lexical retrieval."""

from __future__ import annotations

import hashlib
import re
import sqlite3
import zipfile
from dataclasses import dataclass
from datetime import UTC, datetime
from pathlib import Path

from ato.brain.memory import MemoryItem
from ato.exceptions import MemoryStoreError

MAX_TEXT_DOCUMENT_BYTES = 500_000
MAX_BINARY_DOCUMENT_BYTES = 10_000_000
MAX_EXTRACTED_CHARS = 500_000
MAX_PDF_PAGES = 200
MAX_DOCX_ARCHIVE_BYTES = 20_000_000
MAX_DOCX_ARCHIVE_ENTRIES = 5_000
MAX_DOCUMENT_CHUNKS = 500
MAX_DOCUMENTS = 1_000
CHUNK_CHARS = 1_500
CHUNK_OVERLAP = 200
MAX_SEARCH_CANDIDATES = 2_000
SUPPORTED_SUFFIXES = {
    ".txt",
    ".md",
    ".markdown",
    ".csv",
    ".py",
    ".js",
    ".ts",
    ".tsx",
    ".jsx",
    ".json",
    ".toml",
    ".yaml",
    ".yml",
    ".html",
    ".css",
    ".sql",
    ".sh",
    ".ps1",
    ".pdf",
    ".docx",
}
PROTECTED_PARTS = {".git", ".github", ".venv", "data", "__pycache__"}
WORD_PATTERN = re.compile(r"[a-z0-9]+", re.IGNORECASE)
SECRET_PATTERN = re.compile(
    r"\b(api[_ -]?key|access[_ -]?token|password|secret[_ -]?key)\b\s*[:=]\s*"
    r"[\"']?[a-z0-9_-]{8,}"
    r"|\bsk-[a-z0-9_-]{8,}",
    re.IGNORECASE,
)
SEARCH_STOP_WORDS = {
    "a",
    "an",
    "and",
    "are",
    "for",
    "in",
    "is",
    "of",
    "on",
    "the",
    "to",
    "what",
}


@dataclass(frozen=True, slots=True)
class DocumentRecord:
    id: int
    path: str
    chunks: int


class SqliteKnowledgeStore:
    """Store chunks from explicitly approved workspace documents."""

    def __init__(self, path: Path, workspace_root: Path) -> None:
        self.path = path
        self.workspace_root = workspace_root.resolve()
        self._fts_enabled = False
        self._initialize()

    def ingest(self, relative_path: str) -> DocumentRecord:
        source = self._resolve_source(relative_path)
        try:
            raw_content = source.read_bytes()
        except OSError as exc:
            raise MemoryStoreError("Document could not be read.") from exc
        content = _extract_content(source, raw_content)
        if SECRET_PATTERN.search(content):
            raise MemoryStoreError("Document appears to contain an API key, token, or secret.")
        chunks = _chunk_text(content)
        if not chunks:
            raise MemoryStoreError("Document contains no ingestible text.")
        relative = source.relative_to(self.workspace_root).as_posix()
        digest = hashlib.sha256(raw_content).hexdigest()
        try:
            with self._connect() as connection:
                row = connection.execute(
                    "SELECT id, sha256 FROM documents WHERE path = ?", (relative,)
                ).fetchone()
                if row is None:
                    count = connection.execute("SELECT COUNT(*) FROM documents").fetchone()[0]
                    if count >= MAX_DOCUMENTS:
                        raise MemoryStoreError("Knowledge base has reached its document limit.")
                    cursor = connection.execute(
                        "INSERT INTO documents(path, sha256, updated_at) VALUES (?, ?, ?)",
                        (relative, digest, datetime.now(UTC).isoformat()),
                    )
                    document_id = int(cursor.lastrowid)
                else:
                    document_id = int(row[0])
                    if row[1] == digest:
                        count = connection.execute(
                            "SELECT COUNT(*) FROM chunks WHERE document_id = ?", (document_id,)
                        ).fetchone()[0]
                        return DocumentRecord(document_id, relative, int(count))
                    connection.execute("DELETE FROM chunks WHERE document_id = ?", (document_id,))
                    connection.execute(
                        "UPDATE documents SET sha256 = ?, updated_at = ? WHERE id = ?",
                        (digest, datetime.now(UTC).isoformat(), document_id),
                    )
                connection.executemany(
                    "INSERT INTO chunks(document_id, ordinal, content) VALUES (?, ?, ?)",
                    [(document_id, index, chunk) for index, chunk in enumerate(chunks)],
                )
        except sqlite3.Error as exc:
            raise MemoryStoreError("Document could not be stored in the knowledge base.") from exc
        return DocumentRecord(document_id, relative, len(chunks))

    def list_documents(self) -> tuple[DocumentRecord, ...]:
        try:
            with self._connect() as connection:
                rows = connection.execute(
                    "SELECT d.id, d.path, COUNT(c.id) FROM documents d "
                    "LEFT JOIN chunks c ON c.document_id = d.id GROUP BY d.id ORDER BY d.id DESC"
                ).fetchall()
        except sqlite3.Error as exc:
            raise MemoryStoreError("Knowledge documents could not be listed.") from exc
        return tuple(DocumentRecord(int(row[0]), str(row[1]), int(row[2])) for row in rows)

    def remove_document(self, document_id: int) -> bool:
        if document_id < 1:
            raise ValueError("document_id must be positive.")
        try:
            with self._connect() as connection:
                cursor = connection.execute("DELETE FROM documents WHERE id = ?", (document_id,))
        except sqlite3.Error as exc:
            raise MemoryStoreError("Knowledge document could not be removed.") from exc
        return cursor.rowcount > 0

    def search(self, query: str, limit: int = 5) -> tuple[MemoryItem, ...]:
        if not 1 <= limit <= 20:
            raise ValueError("limit must be between 1 and 20.")
        terms = {
            term
            for term in WORD_PATTERN.findall(query.casefold())
            if term not in SEARCH_STOP_WORDS
        }
        if not terms:
            return ()
        if self._fts_enabled:
            results = self._search_fts(terms, limit)
            if results is not None:
                return results
        return self._search_lexical(terms, limit)

    def _search_fts(self, terms: set[str], limit: int) -> tuple[MemoryItem, ...] | None:
        match_query = " OR ".join(f'"{term}"' for term in sorted(terms))
        try:
            with self._connect() as connection:
                rows = connection.execute(
                    "SELECT c.id, c.content, d.path, c.ordinal FROM chunks_fts f "
                    "JOIN chunks c ON c.id = f.rowid "
                    "JOIN documents d ON d.id = c.document_id "
                    "WHERE chunks_fts MATCH ? ORDER BY bm25(chunks_fts), c.id DESC LIMIT ?",
                    (match_query, limit),
                ).fetchall()
        except sqlite3.OperationalError:
            return None
        return tuple(
            MemoryItem(int(row[0]), str(row[1]), f"knowledge {row[2]}#{row[3]}")
            for row in rows
        )

    def _search_lexical(self, terms: set[str], limit: int) -> tuple[MemoryItem, ...]:
        try:
            with self._connect() as connection:
                rows = connection.execute(
                    "SELECT c.id, c.content, d.path, c.ordinal FROM chunks c "
                    "JOIN documents d ON d.id = c.document_id ORDER BY c.id DESC LIMIT ?",
                    (MAX_SEARCH_CANDIDATES,),
                ).fetchall()
        except sqlite3.Error as exc:
            raise MemoryStoreError("Knowledge base could not be searched.") from exc
        ranked = []
        for row in rows:
            overlap = terms & set(WORD_PATTERN.findall(str(row[1]).casefold()))
            if overlap:
                score = len(overlap) / len(terms)
                item = MemoryItem(int(row[0]), str(row[1]), f"knowledge {row[2]}#{row[3]}")
                ranked.append((score, int(row[0]), item))
        ranked.sort(key=lambda entry: (entry[0], entry[1]), reverse=True)
        return tuple(entry[2] for entry in ranked[:limit])

    def _resolve_source(self, relative_path: str) -> Path:
        requested = Path(relative_path)
        if requested.is_absolute():
            raise MemoryStoreError("Knowledge paths must be workspace-relative.")
        unresolved = self.workspace_root / requested
        if unresolved.is_symlink():
            raise MemoryStoreError("Symbolic-link documents cannot be ingested.")
        source = unresolved.resolve()
        try:
            relative = source.relative_to(self.workspace_root)
        except ValueError as exc:
            raise MemoryStoreError("Document is outside the authorized workspace.") from exc
        if any(part.casefold() in PROTECTED_PARTS for part in relative.parts):
            raise MemoryStoreError("Document is inside a protected workspace directory.")
        if relative.name.casefold() == ".env" or relative.name.casefold().startswith(".env."):
            raise MemoryStoreError("Environment files cannot be ingested.")
        if source.suffix.casefold() not in SUPPORTED_SUFFIXES:
            raise MemoryStoreError("Document type is not supported in this RAG phase.")
        if not source.is_file():
            raise MemoryStoreError("Knowledge path must be an existing file.")
        return source

    def _initialize(self) -> None:
        try:
            self.path.parent.mkdir(parents=True, exist_ok=True)
            with self._connect() as connection:
                connection.execute("PRAGMA foreign_keys = ON")
                connection.execute(
                    "CREATE TABLE IF NOT EXISTS documents (id INTEGER PRIMARY KEY, "
                    "path TEXT NOT NULL UNIQUE, sha256 TEXT NOT NULL, updated_at TEXT NOT NULL)"
                )
                connection.execute(
                    "CREATE TABLE IF NOT EXISTS chunks (id INTEGER PRIMARY KEY, "
                    "document_id INTEGER NOT NULL REFERENCES documents(id) ON DELETE CASCADE, "
                    "ordinal INTEGER NOT NULL, content TEXT NOT NULL)"
                )
                self._fts_enabled = _initialize_fts(connection)
        except (OSError, sqlite3.Error) as exc:
            raise MemoryStoreError("Knowledge base could not be initialized.") from exc

    def _connect(self) -> sqlite3.Connection:
        connection = sqlite3.connect(self.path, timeout=5)
        connection.execute("PRAGMA foreign_keys = ON")
        return connection


def _initialize_fts(connection: sqlite3.Connection) -> bool:
    """Create and synchronize the optional local full-text index."""
    try:
        connection.execute(
            "CREATE TABLE IF NOT EXISTS knowledge_metadata ("
            "key TEXT PRIMARY KEY, value TEXT NOT NULL)"
        )
        connection.execute(
            "CREATE VIRTUAL TABLE IF NOT EXISTS chunks_fts USING fts5("
            "content, content='chunks', content_rowid='id')"
        )
        connection.execute(
            "CREATE TRIGGER IF NOT EXISTS chunks_fts_insert AFTER INSERT ON chunks BEGIN "
            "INSERT INTO chunks_fts(rowid, content) VALUES (new.id, new.content); END"
        )
        connection.execute(
            "CREATE TRIGGER IF NOT EXISTS chunks_fts_delete AFTER DELETE ON chunks BEGIN "
            "INSERT INTO chunks_fts(chunks_fts, rowid, content) "
            "VALUES ('delete', old.id, old.content); END"
        )
        connection.execute(
            "CREATE TRIGGER IF NOT EXISTS chunks_fts_update AFTER UPDATE ON chunks BEGIN "
            "INSERT INTO chunks_fts(chunks_fts, rowid, content) "
            "VALUES ('delete', old.id, old.content); "
            "INSERT INTO chunks_fts(rowid, content) VALUES (new.id, new.content); END"
        )
        version = connection.execute(
            "SELECT value FROM knowledge_metadata WHERE key = 'fts_schema_version'"
        ).fetchone()
        if version is None or version[0] != "1":
            connection.execute("INSERT INTO chunks_fts(chunks_fts) VALUES ('rebuild')")
            connection.execute(
                "INSERT INTO knowledge_metadata(key, value) VALUES ('fts_schema_version', '1') "
                "ON CONFLICT(key) DO UPDATE SET value = excluded.value"
            )
    except sqlite3.OperationalError:
        return False
    return True


def _chunk_text(content: str) -> tuple[str, ...]:
    normalized = content.replace("\r\n", "\n").strip()
    if not normalized:
        return ()
    chunks = []
    start = 0
    while start < len(normalized) and len(chunks) < MAX_DOCUMENT_CHUNKS:
        end = min(len(normalized), start + CHUNK_CHARS)
        if end < len(normalized):
            newline = normalized.rfind("\n", start, end)
            if newline > start + CHUNK_CHARS // 2:
                end = newline
        chunk = normalized[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(normalized):
            break
        start = max(start + 1, end - CHUNK_OVERLAP)
    return tuple(chunks)


def _extract_content(source: Path, raw_content: bytes) -> str:
    suffix = source.suffix.casefold()
    size_limit = (
        MAX_BINARY_DOCUMENT_BYTES if suffix in {".pdf", ".docx"} else MAX_TEXT_DOCUMENT_BYTES
    )
    if len(raw_content) > size_limit:
        raise MemoryStoreError(f"Document exceeds the {size_limit}-byte limit.")
    if suffix == ".pdf":
        return _extract_pdf(source)
    if suffix == ".docx":
        return _extract_docx(source)
    try:
        return raw_content.decode("utf-8")
    except UnicodeDecodeError as exc:
        raise MemoryStoreError("Document must be readable UTF-8 text.") from exc


def _extract_pdf(source: Path) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(source)
        if reader.is_encrypted:
            raise MemoryStoreError("Encrypted PDF documents cannot be ingested.")
        if len(reader.pages) > MAX_PDF_PAGES:
            raise MemoryStoreError(f"PDF exceeds the {MAX_PDF_PAGES}-page limit.")
        sections = []
        for page_number, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                sections.append(f"[PDF page {page_number}]\n{text}")
            if sum(len(section) for section in sections) > MAX_EXTRACTED_CHARS:
                raise MemoryStoreError("Extracted document text exceeds the safety limit.")
        return "\n\n".join(sections)
    except MemoryStoreError:
        raise
    except ImportError as exc:
        raise MemoryStoreError("PDF support requires the pypdf dependency.") from exc
    except Exception as exc:
        raise MemoryStoreError("PDF text could not be extracted.") from exc


def _extract_docx(source: Path) -> str:
    try:
        with zipfile.ZipFile(source) as archive:
            entries = archive.infolist()
            if len(entries) > MAX_DOCX_ARCHIVE_ENTRIES:
                raise MemoryStoreError("DOCX archive contains too many entries.")
            if sum(entry.file_size for entry in entries) > MAX_DOCX_ARCHIVE_BYTES:
                raise MemoryStoreError("DOCX expanded content exceeds the safety limit.")

        from docx import Document

        document = Document(source)
        sections = [paragraph.text.strip() for paragraph in document.paragraphs]
        for table_number, table in enumerate(document.tables, start=1):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows:
                sections.append(f"[DOCX table {table_number}]\n" + "\n".join(rows))
        content = "\n\n".join(section for section in sections if section)
        if len(content) > MAX_EXTRACTED_CHARS:
            raise MemoryStoreError("Extracted document text exceeds the safety limit.")
        return content
    except MemoryStoreError:
        raise
    except ImportError as exc:
        raise MemoryStoreError("DOCX support requires the python-docx dependency.") from exc
    except (OSError, zipfile.BadZipFile, ValueError, KeyError) as exc:
        raise MemoryStoreError("DOCX text could not be extracted.") from exc
