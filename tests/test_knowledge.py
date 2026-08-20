import pytest
from docx import Document
from pypdf import PdfWriter
from reportlab.pdfgen import canvas

from ato.exceptions import MemoryStoreError
from ato.knowledge import SqliteKnowledgeStore


def test_knowledge_ingests_retrieves_updates_and_removes(tmp_path) -> None:
    source = tmp_path / "guide.md"
    source.write_text("# Ghana\nAccra is the capital of Ghana.\n", encoding="utf-8")
    store = SqliteKnowledgeStore(tmp_path / "data" / "knowledge.db", tmp_path)

    document = store.ingest("guide.md")
    duplicate = store.ingest("guide.md")
    results = store.search("What is Ghana's capital?")

    assert duplicate == document
    assert results and "Accra" in results[0].content
    assert "guide.md" in results[0].source
    assert store.list_documents() == (document,)

    source.write_text("# Ghana\nKumasi is a major Ghanaian city.\n", encoding="utf-8")
    updated = store.ingest("guide.md")
    assert updated.id == document.id
    assert store.search("Kumasi")
    assert store.remove_document(document.id) is True
    assert store.search("Kumasi") == ()


@pytest.mark.parametrize("path", [".env", "data/file.txt", "unsupported.exe", "../outside.txt"])
def test_knowledge_rejects_unsafe_or_unsupported_paths(tmp_path, path: str) -> None:
    outside = tmp_path.parent / "outside.txt"
    outside.write_text("outside", encoding="utf-8")
    target = tmp_path / path
    if ".." not in path:
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text("text", encoding="utf-8")
    store = SqliteKnowledgeStore(tmp_path / "knowledge.db", tmp_path)

    with pytest.raises(MemoryStoreError):
        store.ingest(path)


def test_knowledge_rejects_likely_secrets(tmp_path) -> None:
    (tmp_path / "notes.txt").write_text("api_key=do-not-ingest", encoding="utf-8")
    store = SqliteKnowledgeStore(tmp_path / "knowledge.db", tmp_path)

    with pytest.raises(MemoryStoreError, match="secret"):
        store.ingest("notes.txt")


def test_knowledge_ingests_pdf_with_page_markers(tmp_path) -> None:
    source = tmp_path / "guide.pdf"
    pdf = canvas.Canvas(str(source))
    pdf.drawString(72, 720, "Tamale is in northern Ghana.")
    pdf.showPage()
    pdf.drawString(72, 720, "Accra is Ghana's capital.")
    pdf.save()
    store = SqliteKnowledgeStore(tmp_path / "knowledge.db", tmp_path)

    document = store.ingest("guide.pdf")
    results = store.search("Accra capital")

    assert document.chunks == 1
    assert results
    assert "[PDF page 2]" in results[0].content


def test_knowledge_ingests_docx_paragraphs_and_tables(tmp_path) -> None:
    source = tmp_path / "guide.docx"
    document = Document()
    document.add_paragraph("Ato can retrieve approved local knowledge.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Capital"
    table.cell(0, 1).text = "Accra"
    document.save(source)
    store = SqliteKnowledgeStore(tmp_path / "knowledge.db", tmp_path)

    record = store.ingest("guide.docx")
    results = store.search("Capital Accra")

    assert record.chunks == 1
    assert results
    assert "[DOCX table 1]" in results[0].content


def test_knowledge_rejects_encrypted_pdf(tmp_path) -> None:
    source = tmp_path / "private.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    writer.encrypt("password")
    with source.open("wb") as stream:
        writer.write(stream)
    store = SqliteKnowledgeStore(tmp_path / "knowledge.db", tmp_path)

    with pytest.raises(MemoryStoreError, match="Encrypted PDF"):
        store.ingest("private.pdf")


@pytest.mark.parametrize("filename", ["broken.pdf", "broken.docx"])
def test_knowledge_rejects_corrupt_binary_documents(tmp_path, filename: str) -> None:
    (tmp_path / filename).write_bytes(b"not a valid document")
    store = SqliteKnowledgeStore(tmp_path / "knowledge.db", tmp_path)

    with pytest.raises(MemoryStoreError, match="could not be extracted"):
        store.ingest(filename)
